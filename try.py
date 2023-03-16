import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

try:
    os.remove('resume.docx')
except:
    pass


doc=docx.Document()
'''margin : narrow'''
section = doc.sections[0]
section.left_margin = int(914400/4)  # 0.25 inch
section.right_margin = int(914400/4)  # 0.25 inch
section.top_margin = int(914400/4)  # 0.25 inch
section.bottom_margin = int(914400/4)  # 0.25 inch
'''style for section titles'''
secT = doc.styles.add_style('secTitle', WD_STYLE_TYPE.PARAGRAPH)
secT.font.size = Pt(12)
secT.font.name = 'Calibri'
secT.font.bold = True

'''style for section content:'''
secC = doc.styles.add_style('secContent', WD_STYLE_TYPE.PARAGRAPH)
secC.font.size = Pt(10)
secC.font.name = 'Calibri'



def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def titlesection(doc,name,addr,em,mob):
    ns = doc.styles.add_style('name', WD_STYLE_TYPE.PARAGRAPH)
    ns.font.size = Pt(18)
    ns.font.name = 'Calibri'
    ns.font.bold = True
    n = doc.add_paragraph(name.title())
    n.style=ns
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ls=doc.styles.add_style('lines', WD_STYLE_TYPE.PARAGRAPH)
    ls.font.size = Pt(9)
    ls.font.name = 'Calibri'
    line1 = doc.add_paragraph('Address:'+addr+' | '+'Mobile:'+mob+ ' | ')
    line2 = doc.add_paragraph('Email:'+em)
    line1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1.style=ls
    line2.style=ls
    insertHR(line2)
    #empty=doc.add_paragraph()

def careerobj(doc,obj):
    n = doc.add_paragraph('CAREER OBJECTIVE')
    n.style=secT
    insertHR(n)
    p = doc.add_paragraph(obj)
    p.style=secC
    #empty=doc.add_paragraph()

def education(doc,ad):
    n = doc.add_paragraph('EDUCATION')
    n.style=secT
    insertHR(n)
    for ed in ad:
        p = doc.add_paragraph(ed[0]+'\t'+ed[1]+'\t\t'+ed[3]+'\t\t'+ed[2])
        p.style=secC

    #empty=doc.add_paragraph()
    
def skills(doc,line):
    n = doc.add_paragraph('SKILLS')
    n.style=secT
    insertHR(n)
    p = doc.add_paragraph(line)
    p.style=secC
    #empty=doc.add_paragraph()

def projects(doc,all):
    n = doc.add_paragraph('PROJECTS')
    n.style=secT
    insertHR(n)
    for i in all:
        p1 = doc.add_paragraph()
        p1.style=secT

        #https://github.com/python-openxml/python-docx/issues/74
        url=i[2]
        part = p1.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = i[0]
        hyperlink.append(new_run)
        p1._p.append(hyperlink)


        p2=doc.add_paragraph(i[1])
        p2.style=secC
        #empty=doc.add_paragraph()

def workxp(doc,xp):
    n = doc.add_paragraph('EXPERIENCE')
    n.style=secT
    insertHR(n)
    for ed in xp:
        p = doc.add_paragraph(ed[0]+'\t\t'+ed[1])
        p.style=secC

    #empty=doc.add_paragraph()

def section(doc,title,text):
    n = doc.add_paragraph(title.upper())
    n.style=secT
    insertHR(n)

    p = doc.add_paragraph(text)
    p.style=secC

    #empty=doc.add_paragraph()
    




titlesection(doc,'nilesh navalkar','B/505, Sinhagad Building, Rajendra Nagar, Borivali E, Mumbai 400 066','nnavalkar29@gmail.com','9699657488')
careerobj(doc,'To work for an organization which provides me the opportunity to improve my skiils as an engineer and knowledge to grow along with the organization"s objective.')
skills(doc,'c,c++,java,python,html,js')
education(doc,[['B.TECH','9.9sgpa','2021-2022','d. j sanghvi college of engineering'],['B.ed','9.9%','2021-2022','d. j sanghvi college of engineering'],['B.TECH','9.9gpa','2021-2022','d. j sanghvi college of engineering']])
projects(doc,[['title 1','A block-level item flows the text it contains between its left and right edges, adding an additional line each time the text extends beyond its right boundary. For a paragraph, the boundaries are generally the page margins, but they can also be column boundaries if the page is laid out in columns, or cell boundaries if the paragraph occurs inside a table cell.','https://google.com'],['title 1','A block-level item flows the text it contains between its left and right edges, adding an additional line each time the text extends beyond its right boundary. For a paragraph, the boundaries are generally the page margins, but they can also be column boundaries if the page is laid out in columns, or cell boundaries if the paragraph occurs inside a table cell.','https://google.com']])
workxp(doc,[['google intern','2020-2020'],['google intern','2020-2020'],['google intern','2020-2020']])
section(doc,'sample','hobbies:porm,smentai')
doc.save('resume.docx')